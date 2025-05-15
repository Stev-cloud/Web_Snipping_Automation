import pyautogui
import time
from docx import Document
from docx.shared import Inches
import os
from PIL import Image, ImageChops
import shutil

# === CONFIGURATION ===
save_dir = r"C:\Users\HAI\OneDrive - Kumaraguru College of Technology\Documents\scroller"
os.makedirs(save_dir, exist_ok=True)

screenshot_folder = os.path.join(save_dir, "screenshots")
os.makedirs(screenshot_folder, exist_ok=True)

output_file = os.path.join(save_dir, "Website_Document_Capture.docx")

# === CROPPING SETTINGS (based on your screen) ===
crop_left = 120
crop_top = 150
crop_width = 1130
crop_height = 510
crop_box = (crop_left, crop_top, crop_left + crop_width, crop_top + crop_height)

# === INIT ===
doc = Document()

print("🚀 Starting in 5 seconds... Focus the browser window now!")
time.sleep(5)

# Move mouse to center of the page area
center_x = crop_left + crop_width // 2
center_y = crop_top + crop_height // 2
pyautogui.moveTo(center_x, center_y)

page = 1
prev_screenshot = None

while True:
    print(f"📸 Capturing screenshot {page}")
    screenshot_path = os.path.join(screenshot_folder, f"screenshot_{page}.png")

    # Take full screenshot and crop to content area
    full_screenshot = pyautogui.screenshot()
    cropped = full_screenshot.crop(crop_box)
    cropped.save(screenshot_path)

    # Add to Word
    doc.add_paragraph(f"Page {page}")
    doc.add_picture(screenshot_path, width=Inches(6))

    # Compare with previous to detect end
    if prev_screenshot:
        diff = ImageChops.difference(cropped, prev_screenshot)
        if not diff.getbbox():  # No visual change
            print("✅ End of page reached.")
            os.remove(screenshot_path)  # Delete duplicate
            break

    prev_screenshot = cropped
    page += 1

    # Scroll down by simulating mouse wheel
    pyautogui.moveTo(center_x, center_y)
    pyautogui.scroll(-1070)
    time.sleep(1)

# === SAVE RESULTS ===
doc.save(output_file)
print(f"✅ Word document saved at: {output_file}")

# Clean up screenshots
shutil.rmtree(screenshot_folder)
print("🧼 Temporary screenshots deleted.")